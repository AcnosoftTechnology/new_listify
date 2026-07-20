# -*- coding: utf-8 -*-
"""Generate Listify Mobile Push API Word documentation."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = r"E:\xampp\htdocs\listify.asia\docs\Listify_Mobile_Push_Notification_API.docx"


def set_cell_shading(cell, hex_color):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def shade_header_row(row, hex_color="1F4E79"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    # light background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    shade_header_row(hdr)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def main():
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover / Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LISTIFY.ASIA")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Mobile App — Push Notification API")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(47, 84, 150)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run("Developer Documentation\nVersion 1.0  |  July 2026")
    m.font.size = Pt(11)
    m.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    # Overview box
    h = doc.add_heading("1. Overview", level=1)

    doc.add_paragraph(
        "This document describes the push notification APIs for the Listify mobile app. "
        "All mobile APIs run on the API subdomain. The website (www) handles browser push only."
    )

    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Base URL", "https://api.listify.asia/api"],
            ["Authentication", "JWT Bearer token (from login)"],
            ["Content-Type", "application/json"],
            ["Firebase project", "listify-asia-firbase-api"],
            ["Android package", "com.listify.asia"],
            ["Android channel ID", "listify_default"],
        ],
    )

    note = doc.add_paragraph()
    nr = note.add_run("Important: ")
    nr.bold = True
    nr.font.color.rgb = RGBColor(192, 0, 0)
    note.add_run(
        "Do NOT use www.listify.asia for these APIs. "
        "Do NOT use /api/auth/login (deprecated). Use POST /api/login only."
    )

    # Endpoint summary
    doc.add_heading("2. Endpoint Summary", level=1)
    add_table(
        doc,
        ["Method", "Endpoint", "Auth", "Purpose"],
        [
            ["POST", "/api/login", "No", "Get JWT access token"],
            ["POST", "/api/fcm/register", "Bearer JWT", "Register device FCM token"],
            ["POST", "/api/fcm/unregister", "Bearer JWT", "Remove token (logout)"],
            ["DELETE", "/api/fcm/unregister", "Bearer JWT", "Same as unregister"],
            ["GET", "/api/fcm/devices", "Bearer JWT", "List registered devices"],
            ["POST", "/api/fcm/test", "Bearer JWT", "Send test notification"],
        ],
    )

    doc.add_paragraph("Full URLs:")
    add_code(
        doc,
        "https://api.listify.asia/api/login\n"
        "https://api.listify.asia/api/fcm/register\n"
        "https://api.listify.asia/api/fcm/unregister\n"
        "https://api.listify.asia/api/fcm/devices\n"
        "https://api.listify.asia/api/fcm/test",
    )

    # Login
    doc.add_heading("3. Login — Get JWT", level=1)
    doc.add_paragraph("POST /login", style="Heading 3")
    doc.add_paragraph("Request body:")
    add_code(doc, '{\n  "email": "user@example.com",\n  "password": "secret"\n}')
    doc.add_paragraph("Success response (200):")
    add_code(
        doc,
        '{\n'
        '  "status": true,\n'
        '  "token": "<JWT_ACCESS_TOKEN>",\n'
        '  "user": {\n'
        '    "id": 5,\n'
        '    "name": "Abhishek",\n'
        '    "email": "user@example.com",\n'
        '    "role": "2",\n'
        '    "is_agent": 1,\n'
        '    "type": "agent"\n'
        "  }\n"
        "}",
    )
    doc.add_paragraph("Error response (401):")
    add_code(doc, '{\n  "status": false,\n  "message": "Invalid email or password"\n}')
    doc.add_paragraph("App must:")
    for item in [
        "Save token in secure storage.",
        "Send on all protected APIs as: Authorization: Bearer <JWT_ACCESS_TOKEN>",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Register
    doc.add_heading("4. Register FCM Token (Required)", level=1)
    doc.add_paragraph(
        "Call after login, and again whenever Firebase refreshes the device token. "
        "Without this call, the phone will not receive push notifications."
    )
    doc.add_paragraph("POST /fcm/register", style="Heading 3")
    doc.add_paragraph("Header: Authorization: Bearer <JWT>")
    doc.add_paragraph("Request fields:")
    add_table(
        doc,
        ["Field", "Type", "Required", "Description"],
        [
            ["fcm_token", "string", "Yes", "Firebase token from THIS device"],
            ["platform", "string", "Yes", "android | ios | web"],
            ["device_id", "string", "No", "Stable unique device id (recommended)"],
            ["device_label", "string", "No", "e.g. Pixel 8, iPhone 15"],
        ],
    )
    doc.add_paragraph("Example request:")
    add_code(
        doc,
        "{\n"
        '  "fcm_token": "dKx9...:APA91bH...",\n'
        '  "platform": "android",\n'
        '  "device_id": "android-unique-device-id",\n'
        '  "device_label": "Pixel 8"\n'
        "}",
    )
    doc.add_paragraph("Success (200):")
    add_code(
        doc,
        "{\n"
        '  "status": true,\n'
        '  "user_id": 5,\n'
        '  "platform": "android",\n'
        '  "message": "FCM token registered. Device will receive enquiry, chat and order pushes."\n'
        "}",
    )
    doc.add_paragraph("Errors:")
    add_table(
        doc,
        ["Code", "Meaning"],
        [
            ["401", "Missing/invalid JWT — login again"],
            ["422", "Validation failed (fcm_token / platform)"],
            ["503", "Server configuration issue"],
        ],
    )
    warn = doc.add_paragraph()
    wr = warn.add_run("Note: ")
    wr.bold = True
    warn.add_run(
        "FCM token must come from Firebase SDK on the real phone. "
        "Postman cannot generate a real device token."
    )

    # Unregister
    doc.add_heading("5. Unregister Token (Logout)", level=1)
    doc.add_paragraph("POST /fcm/unregister  (also accepts DELETE)")
    doc.add_paragraph("Send one of:")
    add_code(doc, '{ "fcm_token": "dKx9...:APA91bH..." }')
    doc.add_paragraph("or")
    add_code(doc, '{ "device_id": "android-unique-device-id" }')
    doc.add_paragraph("Success:")
    add_code(doc, '{\n  "status": true,\n  "deleted": 1\n}')

    # Devices
    doc.add_heading("6. List Devices (Debug)", level=1)
    doc.add_paragraph("GET /fcm/devices")
    add_code(
        doc,
        "{\n"
        '  "status": true,\n'
        '  "count": 1,\n'
        '  "devices": [\n'
        "    {\n"
        '      "id": 12,\n'
        '      "platform": "android",\n'
        '      "device_id": "android-unique-device-id",\n'
        '      "device_label": "Pixel 8",\n'
        '      "token_preview": "dKx9abc…",\n'
        '      "updated_at": "2026-07-18 12:00:00"\n'
        "    }\n"
        "  ]\n"
        "}",
    )

    # Test
    doc.add_heading("7. Test Push (QA)", level=1)
    doc.add_paragraph("POST /fcm/test — no body required. Sends a test notification to registered devices.")
    doc.add_paragraph("Success:")
    add_code(
        doc,
        "{\n"
        '  "status": true,\n'
        '  "reason": "sent",\n'
        '  "message": "Test push sent — check your phone",\n'
        '  "user_id": 5,\n'
        '  "devices": 1\n'
        "}",
    )
    doc.add_paragraph("Failure examples:")
    add_code(
        doc,
        "{\n"
        '  "status": false,\n'
        '  "reason": "no_tokens",\n'
        '  "message": "No FCM token for this user — call /api/fcm/register first"\n'
        "}",
    )

    # Flow
    doc.add_heading("8. App Integration Flow", level=1)
    steps = [
        "User opens app",
        "POST /login → save JWT",
        "Get FCM token from Firebase SDK (this device)",
        "POST /fcm/register with JWT + fcm_token + platform",
        "On token refresh → POST /fcm/register again",
        "On logout → POST /fcm/unregister",
        "On notification tap → navigate using payload.type",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_heading("Flutter example", level=2)
    add_code(
        doc,
        "// After successful login\n"
        "final jwt = loginResponse['token'];\n"
        "final fcmToken = await FirebaseMessaging.instance.getToken();\n\n"
        "await dio.post(\n"
        "  'https://api.listify.asia/api/fcm/register',\n"
        "  data: {\n"
        "    'fcm_token': fcmToken,\n"
        "    'platform': Platform.isIOS ? 'ios' : 'android',\n"
        "    'device_id': deviceId,\n"
        "    'device_label': deviceModel,\n"
        "  },\n"
        "  options: Options(headers: {\n"
        "    'Authorization': 'Bearer $jwt',\n"
        "    'Content-Type': 'application/json',\n"
        "  }),\n"
        ");\n\n"
        "FirebaseMessaging.instance.onTokenRefresh.listen((newToken) async {\n"
        "  // call /fcm/register again with newToken\n"
        "});",
    )

    # Firebase
    doc.add_heading("9. Firebase Setup (App Side)", level=1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Firebase project", "listify-asia-firbase-api (same as website)"],
            ["Android package", "com.listify.asia"],
            ["Notification channel", "listify_default (exact id)"],
            ["Permissions", "Android 13+: POST_NOTIFICATIONS"],
        ],
    )
    p = doc.add_paragraph()
    p.add_run("Create Android notification channel with ID exactly: ").font.size = Pt(11)
    r = p.add_run("listify_default")
    r.bold = True

    # Payloads
    doc.add_heading("10. Notification Payloads", level=1)
    doc.add_paragraph(
        "After the device token is registered, the server automatically sends pushes for these events."
    )

    doc.add_heading("Enquiry (new appointment)", level=2)
    add_table(
        doc,
        ["Key", "Example"],
        [
            ["type", "enquiry"],
            ["title", "New enquiry received"],
            ["appointment_id", "123"],
            ["listing_id", "45"],
            ["listing_type", "restaurant"],
            ["click_action", "/agent/appointment"],
        ],
    )

    doc.add_heading("Chat message", level=2)
    add_table(
        doc,
        ["Key", "Example"],
        [
            ["type", "chat"],
            ["thread_code", "abc123"],
            ["sender_id", "10"],
            ["url_prefix", "agent or customer"],
            ["click_action", "/agent/messages/{sender_id}/{thread_code}"],
        ],
    )

    doc.add_heading("Shop order", level=2)
    add_table(
        doc,
        ["Key", "Example"],
        [
            ["type", "order"],
            ["order_id", "99"],
            ["listing_id", "45"],
            ["click_action", "/agent/order-manager?order_id=99"],
        ],
    )

    doc.add_heading("Navigation suggestion", level=2)
    add_code(
        doc,
        "switch (data['type']) {\n"
        "  case 'enquiry': open Appointments screen;\n"
        "  case 'chat':    open Chat(thread_code, sender_id);\n"
        "  case 'order':   open Order Manager(order_id);\n"
        "  default:        open Home / Notifications;\n"
        "}",
    )

    # Checklist
    doc.add_heading("11. Developer Checklist", level=1)
    checks = [
        "Firebase project linked (Android + iOS if needed)",
        "Notification permission requested",
        "Android channel listify_default created",
        "Login uses POST /api/login and stores JWT",
        "After login → register FCM token",
        "onTokenRefresh → re-register",
        "Logout → unregister",
        "Handle notification tap by type",
        "QA with /api/fcm/test on a real device",
    ]
    for c in checks:
        doc.add_paragraph(f"☐  {c}")

    # Mistakes
    doc.add_heading("12. Common Mistakes", level=1)
    add_table(
        doc,
        ["Mistake", "Result"],
        [
            ["Using /api/auth/login", "404 Not Found"],
            ["Using www.listify.asia for FCM APIs", "Wrong / missing routes"],
            ["Missing Authorization: Bearer", "401 Unauthenticated"],
            ["Fake / Postman token", "No real phone notification"],
            ["Skipping /fcm/register", "No pushes on that device"],
            ["Wrong Firebase project / package", "Invalid token / no delivery"],
        ],
    )

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Listify.asia API Team  •  Confidential — for app development use only  •  July 2026")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUT)
    # also copy next to md
    alt = r"E:\xampp\htdocs\listify.asia\Listify_Mobile_Push_Notification_API.docx"
    doc.save(alt)
    alt2 = r"E:\xampp\htdocs\api.listify.aisa\Listify_Mobile_Push_Notification_API.docx"
    doc.save(alt2)
    print(OUT)
    print(alt)
    print(alt2)


if __name__ == "__main__":
    main()
